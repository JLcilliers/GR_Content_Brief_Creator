"""
Document Formatter
Creates formatted Word documents from content brief data.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from typing import Dict
import os
from datetime import datetime


class DocumentFormatter:
    """Formats content briefs into Word documents with consistent styling."""
    
    def __init__(self):
        self.font_name = "Calibri"
        self.body_size = 11
        self.heading_size = 12
        
        # Professional color scheme
        self.header_bg = RGBColor(0, 32, 96)  # Dark blue
        self.header_text = RGBColor(255, 255, 255)  # White
        self.section_bg = RGBColor(68, 114, 196)  # Medium blue
        self.section_text = RGBColor(255, 255, 255)  # White
        self.content_bg = RGBColor(242, 242, 242)  # Light gray
        self.content_text = RGBColor(0, 0, 0)  # Black
    
    def create_brief_document(self, brief_data: Dict, output_dir: str = "output_briefs") -> str:
        """Create a formatted Word document from brief data."""
        
        doc = Document()
        
        # Set default font for the document
        style = doc.styles['Normal']
        font = style.font
        font.name = self.font_name
        font.size = Pt(self.body_size)
        
        # Header
        self._add_header(doc, brief_data)
        
        # Add each section
        self._add_section(doc, "1. Page Type Identification", brief_data.get("page_type", ""))
        self._add_section(doc, "2. Page Title", brief_data.get("page_title", ""))
        self._add_section(doc, "3. Meta Description", brief_data.get("meta_description", ""))
        self._add_section(doc, "4. Target URL", brief_data.get("target_url", ""))
        self._add_section(doc, "5. H1 Heading", brief_data.get("h1", ""))
        self._add_section(doc, "6. Summary Bullets", brief_data.get("summary_bullets", ""))
        self._add_section(doc, "7. Internal Linking", brief_data.get("internal_links", ""))
        self._add_section(doc, "8. Audience Definition", brief_data.get("audience", ""))
        self._add_section(doc, "9. CTA / Path", brief_data.get("cta", ""))
        self._add_section(doc, "10. Restrictions", brief_data.get("restrictions", ""))
        self._add_section(doc, "11. Requirements", brief_data.get("requirements", ""))
        self._add_section(doc, "12. Suggested Headings & Key Points (+ FAQ)", brief_data.get("headings_faq", ""))
        
        # Save the document
        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        client_name = brief_data.get("client_name", "Client").replace(" ", "_")
        topic = brief_data.get("topic", "Topic").replace(" ", "_")[:30]  # Limit length
        
        filename = f"{client_name}_{topic}_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)
        
        doc.save(filepath)
        
        return filepath
    
    def _add_header(self, doc: Document, brief_data: Dict):
        """Add professional table-based header with client information."""
        
        # Main title with colored background
        title_table = doc.add_table(rows=1, cols=1)
        title_table.style = 'Table Grid'
        title_cell = title_table.rows[0].cells[0]
        title_cell.text = f"{brief_data.get('client_name', '')} - {brief_data.get('topic', '')} - Content Brief"
        
        # Style title cell
        title_para = title_cell.paragraphs[0]
        title_run = title_para.runs[0]
        title_run.font.name = self.font_name
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = self.header_text
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Set title cell background color
        self._set_cell_background(title_cell, self.header_bg)
        
        doc.add_paragraph()  # Spacing
        
        # Info table with metadata
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        # Set column widths
        info_table.columns[0].width = Inches(1.5)
        info_table.columns[1].width = Inches(5.0)
        
        # Row 1: Site
        self._add_info_row(info_table.rows[0], "Site:", brief_data.get('site', ''))
        
        # Row 2: Primary Keyword
        self._add_info_row(info_table.rows[1], "Primary Keyword:", brief_data.get('primary_kw', ''))
        
        # Row 3: Secondary Keywords
        self._add_info_row(info_table.rows[2], "Secondary Keywords:", brief_data.get('secondary_kws', ''))
        
        # Row 4: Date Generated
        self._add_info_row(info_table.rows[3], "Date Generated:", datetime.now().strftime("%d %B %Y"))
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_section(self, doc: Document, section_title: str, content: str):
        """Add a professional section with table-based layout."""
        
        # Create section table
        section_table = doc.add_table(rows=2, cols=1)
        section_table.style = 'Table Grid'
        
        # Title row with colored background
        title_cell = section_table.rows[0].cells[0]
        title_cell.text = section_title
        
        # Style title
        title_para = title_cell.paragraphs[0]
        title_run = title_para.runs[0]
        title_run.font.name = self.font_name
        title_run.font.size = Pt(self.heading_size)
        title_run.font.bold = True
        title_run.font.color.rgb = self.section_text
        
        # Set title cell background
        self._set_cell_background(title_cell, self.section_bg)
        
        # Content row with light background
        content_cell = section_table.rows[1].cells[0]
        
        # Set content cell background
        self._set_cell_background(content_cell, self.content_bg)
        
        # Add content with proper formatting
        content_cell.text = ''  # Clear default text
        content_lines = content.split('\n')
        
        for i, line in enumerate(content_lines):
            if i > 0:  # Add paragraph for subsequent lines
                content_cell.add_paragraph()
            
            if line.strip():
                para = content_cell.paragraphs[-1]
                run = para.add_run(line)
                run.font.name = self.font_name
                run.font.size = Pt(self.body_size)
                run.font.color.rgb = self.content_text
        
        # Add spacing after section
        doc.add_paragraph()
    
    def set_font(self, font_name: str):
        """Change the font name."""
        self.font_name = font_name
    
    def set_body_size(self, size: int):
        """Change the body font size."""
        self.body_size = size
    
    def set_heading_size(self, size: int):
        """Change the heading font size."""
        self.heading_size = size
    
    def _set_cell_background(self, cell, color: RGBColor):
        """Set background color for a table cell."""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(color.r, color.g, color.b))
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def _add_info_row(self, row, label: str, value: str):
        """Add a formatted info row to the header table."""
        # Label cell
        label_cell = row.cells[0]
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.font.name = self.font_name
        label_run.font.size = Pt(self.body_size)
        label_run.font.bold = True
        self._set_cell_background(label_cell, self.content_bg)
        
        # Value cell
        value_cell = row.cells[1]
        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run(value)
        value_run.font.name = self.font_name
        value_run.font.size = Pt(self.body_size)
